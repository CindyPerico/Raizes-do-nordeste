"""Gera a versao editavel (.docx) do trabalho academico.

Reaproveita o conteudo de gerar_pdf.py: as funcoes de montagem sao substituidas
por versoes que produzem blocos neutros, que aqui sao renderizados com python-docx.

Uso:
    python3 docs/pdf/gerar_docx.py --aluno "Cindy Perico" --ru 4906176 \
        --saida "docs/pdf/4906176_Projeto_Back_End.docx"
"""

import argparse
import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import gerar_pdf as g  # noqa: E402

BASE = g.BASE
IMG = g.IMG


def _blocos_conteudo(repo):
    """Executa gerar_pdf.conteudo com as funcoes de montagem substituidas."""
    caption_ids = {id(g.CAPTION)}
    ref_ids = {id(g.REF)}

    g.h = lambda texto, nivel=0: ("h", nivel, texto)
    g.p = lambda texto, recuo=True: ("p", texto, recuo)
    g.itens = lambda linhas: [("li", x) for x in linhas]
    g.code = lambda texto: ("code", texto.strip("\n"))
    g.tabela = lambda cabecalho, linhas, larguras: ("tabela", cabecalho, linhas, larguras)
    g.figura = lambda arquivo, legenda, largura_cm=15.0, altura_max_cm=19.0: [
        ("img", os.path.join(IMG, arquivo), largura_cm), ("cap", legenda)]
    g.PageBreak = lambda: ("pb",)
    g.Spacer = lambda *a, **k: ("sp",)
    g.Paragraph = lambda texto, estilo=None: (
        "cap" if id(estilo) in caption_ids else "ref" if id(estilo) in ref_ids else "p",
        texto) if id(estilo) in caption_ids | ref_ids else ("p", texto, True)

    return g.conteudo(repo)


TAGS = re.compile(r"(<b>|</b>|<i>|</i>|<br\s*/?>)")


def _escreve(par, texto):
    negrito = italico = False
    for parte in TAGS.split(texto):
        if parte == "<b>":
            negrito = True
        elif parte == "</b>":
            negrito = False
        elif parte == "<i>":
            italico = True
        elif parte == "</i>":
            italico = False
        elif parte.startswith("<br"):
            par.add_run().add_break()
        elif parte:
            run = par.add_run(parte.replace("&amp;", "&").replace("&lt;", "<")
                              .replace("&gt;", ">").replace("&nbsp;", " "))
            run.bold = negrito
            run.italic = italico


def _paragrafo(doc, texto, estilo="Corpo", alinhamento=None, recuo=True):
    par = doc.add_paragraph(style=estilo)
    if alinhamento is not None:
        par.alignment = alinhamento
    if estilo == "Corpo" and not recuo:
        par.paragraph_format.first_line_indent = Cm(0)
    _escreve(par, texto)
    return par


def _fonte(estilo, nome):
    rpr = estilo.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for atributo in ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme"):
        if rfonts.get(qn(atributo)) is not None:
            del rfonts.attrib[qn(atributo)]
    for atributo in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(atributo), nome)


def _estilos(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    _fonte(normal, "Times New Roman")
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    corpo = doc.styles.add_style("Corpo", 1)
    corpo.base_style = normal
    corpo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    corpo.paragraph_format.line_spacing = 1.5
    corpo.paragraph_format.first_line_indent = Cm(1.25)
    corpo.paragraph_format.space_after = Pt(0)

    for nivel, tamanho in ((1, 12), (2, 12), (3, 12)):
        estilo = doc.styles[f"Heading {nivel}"]
        estilo.font.name = "Times New Roman"
        estilo.font.size = Pt(tamanho)
        estilo.font.bold = True
        estilo.font.italic = nivel == 3
        estilo.font.color.rgb = RGBColor(0, 0, 0)
        _fonte(estilo, "Times New Roman")
        estilo.paragraph_format.space_before = Pt(18 if nivel == 1 else 12)
        estilo.paragraph_format.space_after = Pt(6)
        estilo.paragraph_format.keep_with_next = True

    legenda = doc.styles.add_style("Legenda", 1)
    legenda.base_style = normal
    legenda.font.size = Pt(10)
    legenda.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    legenda.paragraph_format.space_after = Pt(12)

    codigo = doc.styles.add_style("Codigo", 1)
    codigo.base_style = normal
    codigo.font.name = "Courier New"
    _fonte(codigo, "Courier New")
    codigo.font.size = Pt(9)
    codigo.paragraph_format.line_spacing = 1.0
    codigo.paragraph_format.left_indent = Cm(0.5)
    codigo.paragraph_format.space_after = Pt(10)

    lista = doc.styles.add_style("Lista", 1)
    lista.base_style = normal
    lista.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    lista.paragraph_format.line_spacing = 1.5
    lista.paragraph_format.left_indent = Cm(1.25)
    lista.paragraph_format.space_after = Pt(0)

    celula = doc.styles.add_style("Celula", 1)
    celula.base_style = normal
    celula.font.size = Pt(9)
    celula.paragraph_format.line_spacing = 1.0
    celula.paragraph_format.space_after = Pt(0)

    referencia = doc.styles.add_style("Referencia", 1)
    referencia.base_style = normal
    referencia.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    referencia.paragraph_format.space_after = Pt(12)


def _numeracao(section):
    """Numero de pagina no canto superior direito, a partir da primeira pagina textual."""
    par = section.header.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    par.style = section.header.paragraphs[0].style
    campo = OxmlElement("w:fldSimple")
    campo.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    texto = OxmlElement("w:t")
    texto.text = "5"
    run.append(texto)
    campo.append(run)
    par._p.append(campo)
    for run in par.runs:
        run.font.size = Pt(10)


def _sumario(doc):
    par = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    aviso = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = ("Clique com o botao direito e escolha \u201cAtualizar campo\u201d "
              "para gerar o sumario.")
    aviso.append(t)
    fld.append(aviso)
    par._p.append(fld)


def _capa(doc, aluno, ru, ano):
    for texto in ("CENTRO UNIVERSITÁRIO INTERNACIONAL UNINTER",
                  "PROJETO MULTIDISCIPLINAR — TRILHA BACK-END"):
        _paragrafo(doc, texto, "Normal", WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(11):
        doc.add_paragraph()
    _paragrafo(doc, f"<b>{aluno.upper()}</b>", "Normal", WD_ALIGN_PARAGRAPH.CENTER)
    _paragrafo(doc, f"RU: {ru}", "Normal", WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(6):
        doc.add_paragraph()
    _paragrafo(doc, "<b>REDE DE LANCHONETES “RAÍZES DO NORDESTE”:<br/>"
                    "PROJETO E IMPLEMENTAÇÃO DE UMA API REST MULTICANAL</b>",
               "Normal", WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(19):
        doc.add_paragraph()
    _paragrafo(doc, "CURITIBA", "Normal", WD_ALIGN_PARAGRAPH.CENTER)
    _paragrafo(doc, str(ano), "Normal", WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    _paragrafo(doc, f"<b>{aluno.upper()}</b>", "Normal", WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(11):
        doc.add_paragraph()
    _paragrafo(doc, "<b>REDE DE LANCHONETES “RAÍZES DO NORDESTE”:<br/>"
                    "PROJETO E IMPLEMENTAÇÃO DE UMA API REST MULTICANAL</b>",
               "Normal", WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(5):
        doc.add_paragraph()
    natureza = doc.add_paragraph()
    natureza.paragraph_format.left_indent = Cm(8)
    natureza.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _escreve(natureza, "Trabalho apresentado à disciplina de Projeto Multidisciplinar — "
                       "Trilha Back-End do Centro Universitário Internacional UNINTER, como "
                       "requisito parcial para avaliação da atividade prática.<br/><br/>"
                       "Prof. Me. Luciane Yanase Kanashiro.")
    for _ in range(17):
        doc.add_paragraph()
    _paragrafo(doc, "CURITIBA", "Normal", WD_ALIGN_PARAGRAPH.CENTER)
    _paragrafo(doc, str(ano), "Normal", WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    _paragrafo(doc, "<b>SUMÁRIO</b>", "Normal", WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    _sumario(doc)


def _tabela(doc, cabecalho, linhas, larguras):
    tabela = doc.add_table(rows=1, cols=len(cabecalho))
    tabela.style = "Table Grid"
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = False
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tabela._tbl.tblPr.append(layout)
    for indice, celula in enumerate(tabela.rows[0].cells):
        celula.paragraphs[0].style = doc.styles["Celula"]
        _escreve(celula.paragraphs[0], f"<b>{cabecalho[indice]}</b>")
    for linha in linhas:
        celulas = tabela.add_row().cells
        for indice, valor in enumerate(linha):
            celulas[indice].paragraphs[0].style = doc.styles["Celula"]
            _escreve(celulas[indice].paragraphs[0], str(valor))
    escala = min(1.0, 16.0 / sum(larguras))
    for linha in tabela.rows:
        for indice, celula in enumerate(linha.cells):
            celula.width = Cm(larguras[indice] * escala)
    doc.add_paragraph()


def _renderiza(doc, blocos):
    for bloco in blocos:
        tipo = bloco[0]
        if tipo == "h":
            _, nivel, texto = bloco
            par = doc.add_paragraph(style=f"Heading {nivel + 1}")
            _escreve(par, texto)
            for run in par.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.bold = True
                run.italic = nivel == 2
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif tipo == "p":
            _paragrafo(doc, bloco[1], "Corpo", recuo=bloco[2] if len(bloco) > 2 else True)
        elif tipo == "li":
            _paragrafo(doc, f"• {bloco[1]}", "Lista")
        elif tipo == "code":
            for linha in bloco[1].split("\n"):
                par = doc.add_paragraph(style="Codigo")
                _escreve(par, linha or " ")
        elif tipo == "cap":
            _paragrafo(doc, bloco[1], "Legenda")
        elif tipo == "ref":
            _paragrafo(doc, bloco[1], "Referencia")
        elif tipo == "tabela":
            _tabela(doc, bloco[1], bloco[2], bloco[3])
        elif tipo == "img":
            par = doc.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            par.add_run().add_picture(bloco[1], width=Cm(min(bloco[2], 16.0)))
        elif tipo == "pb":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        elif tipo == "sp":
            doc.add_paragraph()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--aluno", default="Cindy Perico")
    parser.add_argument("--ru", default="0000000")
    parser.add_argument("--ano", default="2026")
    parser.add_argument("--repo", default="https://github.com/CindyPerico/trabalhoback-end")
    parser.add_argument("--saida", default=os.path.join(BASE, "docs", "pdf",
                                                        "Projeto_Back_End.docx"))
    args = parser.parse_args()

    blocos = _blocos_conteudo(args.repo)

    doc = Document()
    secao = doc.sections[0]
    secao.page_width, secao.page_height = Cm(21), Cm(29.7)
    secao.left_margin, secao.top_margin = Cm(3), Cm(3)
    secao.right_margin, secao.bottom_margin = Cm(2), Cm(2)
    _estilos(doc)
    _capa(doc, args.aluno, args.ru, args.ano)

    corpo = doc.add_section(WD_SECTION.NEW_PAGE)
    corpo.page_width, corpo.page_height = Cm(21), Cm(29.7)
    corpo.left_margin, corpo.top_margin = Cm(3), Cm(3)
    corpo.right_margin, corpo.bottom_margin = Cm(2), Cm(2)
    corpo.header.is_linked_to_previous = False
    _numeracao(corpo)

    _renderiza(doc, blocos)
    doc.save(args.saida)
    print(f"DOCX gerado em {args.saida}")


if __name__ == "__main__":
    main()
